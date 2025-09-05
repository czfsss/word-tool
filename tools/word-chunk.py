from collections.abc import Generator
from typing import Any
import tempfile
import os
import re
import json
from dify_plugin.file.file import File
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tools.utils.logger_utils import get_logger


class WordChunkTool(Tool):
    # 获取当前模块的日志记录器
    logger = get_logger(__name__)

    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        # 获取上传的word文件
        word_content: File = tool_parameters.get("word_content")
        chunk_num: int = tool_parameters.get("chunk_num")
        docx_type: str = tool_parameters.get("docx_type")

        if not word_content:
            self.logger.error("未提供Word文件")
            yield self.create_text_message("请提供word文件")
            return
        # 检查文件类型
        if not isinstance(word_content, File):
            self.logger.error("无效的文件格式，期望File对象")
            yield self.create_text_message("无效的文件格式，期望File对象")
            return

        self.logger.info(
            f"开始处理Word分块，文件名: {word_content.filename if word_content.filename else '未知'}，目标分块数: {chunk_num}"
        )

        try:
            # 创建临时文件保存上传的word文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                # 获取文件内容（字节）并写入临时文件
                temp_file.write(word_content.blob)
                temp_file_path = temp_file.name

            # 调用分段函数
            self.logger.info("开始执行智能分段")
            chunks = self.smart_chunk_paragraphs(temp_file_path, doc_type=docx_type)
            self.logger.info(f"初始分段完成，共生成 {len(chunks)} 个段落")

            # 限制分段个数不超过30个
            chunks = self.limit_chunks_to_max(chunks, max_chunks=chunk_num)
            self.logger.info(f"分段数量限制完成，最终段落数: {len(chunks)}")

            # 清理临时文件
            os.unlink(temp_file_path)

            # 返回分段结果
            result = {str(i + 1): chunk for i, chunk in enumerate(chunks)}
            self.logger.info(f"Word分块处理完成，成功生成 {len(result)} 个分块")
            yield self.create_json_message(result)

        except Exception as e:
            self.logger.exception("处理Word文件时发生异常")
            yield self.create_text_message(f"处理word文件时出错: {str(e)}")

    def smart_chunk_paragraphs(self, doc_path, min_length=1000, doc_type=None):
        """
        智能合并短段落，生成有意义的文本块，特别优化了合同和制度文件的处理。
        参数:
            doc_path: Word文档路径
            min_length: 被认为是有独立意义的最小段落长度（字符数）
            doc_type: 文档类型，可选"general"（通用）、"contract"（合同）、"policy"（制度文件）
        返回:
            一个包含合并后文本块的列表
        """
        doc = Document(doc_path)
        chunks = []  # 最终返回的块列表
        current_chunk = []  # 当前正在构建的块（由多个段落组成）
        consecutive_title_count = 0  # 连续标题计数器
        current_section_level = 0  # 当前章节层级（用于制度文件）
        current_clause_level = 0  # 当前条款层级（用于合同文件）
        
        # 根据文档类型设置特殊参数
        if doc_type == "contract":
            # 合同文件通常条款较短，可以适当降低min_length
            min_length = min(min_length, 800)
            # 合同文件的特殊段落标记
            special_markers = ["鉴于", "第一条", "甲方：", "乙方：", "合同编号：", "签订日期："]
        elif doc_type == "policy":
            # 制度文件通常章节较长，可以适当增加min_length
            min_length = max(min_length, 1200)
            # 制度文件的特殊段落标记
            special_markers = ["第一章", "第一条", "总则", "附则", "附件"]
        else:
            special_markers = []

        # 处理所有段落和表格
        elements = self._get_document_elements(doc)
        
        for element in elements:
            element_type = element["type"]
            text = element["text"].strip()
            paragraph = element.get("paragraph")
            
            # 跳过完全空的段落（通常是格式性的换行）
            if not text:
                continue

            # 使用增强的标题判断函数，传入文档类型
            is_heading = False
            if element_type == "paragraph" and paragraph:
                is_heading = self.is_title(paragraph, doc_type=doc_type)
            

            
            # 检测特殊段落（如合同签署区、制度文件附件等）
            is_special = False
            if element_type == "paragraph":
                if doc_type == "contract" and any(marker in text for marker in ["签署页", "双方签字", "附件"]):
                    is_special = True
                elif doc_type == "policy" and any(marker in text for marker in ["附件", "附表", "附图"]):
                    is_special = True



            # 情况1：遇到一个标题
            if is_heading:
                
                consecutive_title_count += 1  # 增加连续标题计数
                
                # 更新章节/条款层级
                if doc_type == "policy":
                    if "章" in text:
                        current_section_level = 1
                    elif "节" in text:
                        current_section_level = 2
                    elif "条" in text:
                        current_section_level = 3
                    elif "款" in text:
                        current_section_level = 4
                    elif "项" in text:
                        current_section_level = 5
                elif doc_type == "contract":
                    if "条" in text:
                        # 提取条款号
                        match = re.search(r"第([一二三四五六七八九十百千万\d]+)条", text)
                        if match:
                            current_clause_level = int(match.group(1)) if match.group(1).isdigit() else len(match.group(1))
                
                # 检查是否为数字标题（如1.1）
                is_numbered_heading = re.match(r"^\d+(\.\d+)?\s*", text)
                
                # 如果当前块不为空，且不是连续标题的情况，则保存当前块
                if current_chunk and consecutive_title_count == 1:
                    # 检查当前块是否以数字标题开头，且当前标题是它的子标题
                    if current_chunk and is_numbered_heading:
                        first_line = current_chunk[0].strip()
                        # 检查第一行是否为数字标题（如1.1）
                        first_is_numbered = re.match(r"^(\d+(\.\d+)?)\s*", first_line)
                        if first_is_numbered:
                            # 获取当前标题的数字部分
                            current_num = text.split('.')[0]
                            first_num = first_line.split('.')[0]
                            # 如果当前标题是第一行标题的子标题（如1.1.1是1.1的子标题）
                            if current_num == first_num and text.count('.') > first_line.count('.'):
                                # 将子标题合并到当前块，而不是创建新块
                                current_chunk.append(text)
                                consecutive_title_count = 0  # 重置计数器
                                continue
                    
                    chunks.append("\n".join(current_chunk))
                    current_chunk = [text]  # 新块以标题开始
                # 如果是连续标题，则添加到当前块
                else:
                    current_chunk.append(text)

            # 情况2：遇到非标题段落
            else:
                # 处理表格内容 - 确保表格与最近的标题在同一个块中
                # 表格通常有表格标题，必须与最近的标题在同一个块中
                if element_type == "table":
                    # 如果当前块不为空（通常包含最近的标题），直接将表格添加到当前块
                    if current_chunk:
                        current_chunk.append(text)
                    # 如果当前块为空，但有之前的块（最后一个块应该包含最近的标题），将表格添加到最后一个块中
                    elif chunks:
                        last_chunk = chunks[-1]
                        chunks[-1] = last_chunk + "\n" + text
                    # 如果既没有当前块也没有之前的块（表格是文档的第一个元素），则创建新块
                    else:
                        current_chunk.append(text)
                    consecutive_title_count = 0  # 重置计数器
                    continue
                
                # 如果之前有连续标题，则将非标题段落与之前的标题合并
                if consecutive_title_count > 0:
                    current_chunk.append(text)
                    consecutive_title_count = 0  # 重置计数器
                # 情况3：当前段落很短，且当前块不为空 -> 合并到当前块
                elif len(text) < min_length and current_chunk:
                    current_chunk.append(text)
                # 情况4：当前段落很长 -> 它自己可以成为一个有意义的块
                elif len(text) >= min_length and not current_chunk:
                    chunks.append(text)
                # 情况5：当前段落很长，但当前块已有内容 -> 结束当前块，并以此段落开始新块
                elif len(text) >= min_length and current_chunk:
                    chunks.append("\n".join(current_chunk))
                    current_chunk = [text]
                # 情况6：当前段落很短，且当前块为空 -> 开始一个新块（希望后续段落能合并进来）
                else:
                    current_chunk.append(text)

        # 循环结束后，处理剩余的块
        if current_chunk:
            chunks.append("\n".join(current_chunk))
            


        return chunks
    
    def limit_chunks_to_max(self, chunks, max_chunks=30):
        """
        限制分段个数不超过指定数量，如果超过则按数学方法合并相邻段落。

        合并策略：
        1. 计算 商 = 段落总数 // max_chunks，余数 = 段落总数 % max_chunks
        2. 如果余数为0，每「商」个段落合并成1个
        3. 如果余数不为0，前「余数」组每「商+1」个段落合并成1个，剩余的每「商」个段落合并成1个

        参数:
            chunks: 原始分段列表
            max_chunks: 最大分段个数，默认30
        返回:
            合并后的分段列表（最多 max_chunks 个）
        """
        if len(chunks) <= max_chunks:
            return chunks

        total_chunks = len(chunks)
        quotient = total_chunks // max_chunks  # 商
        remainder = total_chunks % max_chunks  # 余数

        result_chunks = []
        chunk_index = 0

        # 判断是否能整除
        if total_chunks % max_chunks == 0:
            # 能整除：每 quotient 个段落合并成1个
            for i in range(max_chunks):
                merged_chunk_parts = []
                for j in range(quotient):
                    if chunk_index < total_chunks:
                        merged_chunk_parts.append(chunks[chunk_index])
                        chunk_index += 1

                if merged_chunk_parts:
                    result_chunks.append("\n".join(merged_chunk_parts))
        else:
            # 不能整除：前 remainder 组每组 quotient+1 个段落，剩余的每组 quotient 个段落

            # 处理前 remainder 组，每组 quotient+1 个段落
            for i in range(remainder):
                # 合并 quotient+1 个段落
                merged_chunk_parts = []
                for j in range(quotient + 1):
                    if chunk_index < total_chunks:
                        merged_chunk_parts.append(chunks[chunk_index])
                        chunk_index += 1

                if merged_chunk_parts:
                    result_chunks.append("\n".join(merged_chunk_parts))

            # 处理剩余的组，每组 quotient 个段落
            remaining_groups = max_chunks - remainder
            for i in range(remaining_groups):
                # 合并 quotient 个段落
                merged_chunk_parts = []
                for j in range(quotient):
                    if chunk_index < total_chunks:
                        merged_chunk_parts.append(chunks[chunk_index])
                        chunk_index += 1

                if merged_chunk_parts:
                    result_chunks.append("\n".join(merged_chunk_parts))

        return result_chunks

    def _get_document_elements(self, doc):
        """
        提取文档中的所有段落和表格，按照它们在文档中出现的顺序返回。
        
        参数:
            doc: Document对象
            
        返回:
            一个包含所有段落和表格的列表，每个元素是一个字典，包含type和text字段
            type可以是"paragraph"或"table"，text是段落或表格的文本内容
        """
        elements = []
        
        # 遍历文档中的所有段落
        for paragraph in doc.paragraphs:
            elements.append({
                "type": "paragraph",
                "text": paragraph.text,
                "paragraph": paragraph
            })
        
        # 遍历文档中的所有表格
        for table in doc.tables:
            table_text = ""
            # 处理表格中的每一行
            for row in table.rows:
                row_text = []
                # 处理行中的每个单元格
                for cell in row.cells:
                    # 处理单元格中的每个段落
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        cell_text.append(paragraph.text.strip())
                    # 将单元格中的所有段落文本合并
                    row_text.append(" | ".join(cell_text))
                # 将行中的所有单元格文本合并，并用换行符分隔
                table_text += "\t".join(row_text) + "\n"
            
            elements.append({
                "type": "table",
                "text": table_text.strip()
            })
        
        return elements

    def is_title(self, paragraph, median_size=10, doc_type=None):
        """
        增强版标题识别函数，特别优化了合同和制度文件的标题识别
        参数:
            paragraph: Word文档段落对象
            median_size: 中位数字体大小，用于动态字体大小检测
            doc_position: 段落在文档中的位置（如"doc_start", "section_start"等）
            doc_type: 文档类型，可选"general"（通用）、"contract"（合同）、"policy"（制度文件）
        返回:
            bool: 是否为标题
        """
        text = paragraph.text.strip()
        
        # 改进排除规则（增加长度判断）
        if any(k in text for k in ["图", "表", "注："]) and len(text) < 15:
            return False
            
        # 排除以特殊符号开头的段落（这些通常不是标题）
        if re.match(r"^[（(][^）)]*[）)]|^--+|^—+|^==+|^__+", text):
            return False

        # 1. 增强样式检测（增加常见样式别名）
        if hasattr(paragraph, "style") and paragraph.style:
            style_name = paragraph.style.name.lower()
            style_keywords = ["heading", "title", "标题", "chapter", "header", "titulo"]
            if any(kw in style_name for kw in style_keywords):
                return True

        # 2. 合同/制度文件特有的标题模式
        if doc_type in ["contract", "policy"]:
            # 合同特有模式
            contract_patterns = [
                r"^第[一二三四五六七八九十百千万\d]+条[：: ]",  # 第X条
                r"^(甲方|乙方|丙方|丁方)[：: ]",              # 合同方
                r"^合同编号[：: ]",                          # 合同编号
                r"^签订日期[：: ]",                          # 签订日期
                r"^签订地点[：: ]",                          # 签订地点
                r"^鉴于[：: ]",                              # 鉴于条款
                r"^附件[一二三四五六七八九十百千万\d]*[：: ]",  # 附件
            ]
            
            # 制度文件特有模式
            policy_patterns = [
                r"^第[一二三四五六七八九十百千万\d]+章[：: ]",   # 第X章
                r"^第[一二三四五六七八九十百千万\d]+节[：: ]",   # 第X节
                r"^第[一二三四五六七八九十百千万\d]+款[：: ]",   # 第X款
                r"^第[一二三四五六七八九十百千万\d]+项[：: ]",   # 第X项
                r"^[（(][一二三四五六七八九十百千万\d]+[）)][：: ]",  # (X)
                r"^总则[：: ]",                               # 总则
                r"^附则[：: ]",                               # 附则
            ]
            
            # 根据文档类型选择模式
            patterns = contract_patterns if doc_type == "contract" else policy_patterns
            if any(re.match(p, text) for p in patterns):
                return True
        
        # 3. 居中对齐检测（新增）
        if hasattr(paragraph, "paragraph_format") and paragraph.paragraph_format.alignment:
            if paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                # 居中对齐+合理长度 = 强标题信号
                if 3 <= len(text) <= 50:
                    return True

        # 4. 重构正则表达式（修改：不将1.1.1这样的小级标题识别为标题）
        patterns = [
            r"^\d+\.\d*\s*",              # 一级或二级标题 1. 或 1.1
            r"^[IVXLCDM]+\.\s",            # 罗马数字 I.
            r"^第[\u4e00-\u9fa5\d]+[章节条]", # 第X章/节/条
            r"^[一二三四五六七八九十百千万]+[、.]\s?", # 中文序号
            r"^[A-Z]{2,}[A-Z\s]*\b",       # 全大写短语（至少2字母）
            r"^附 ?录\s?[A-Z]：?"            # 附录类标题
        ]
        if any(re.match(p, text) for p in patterns):
            # 检查是否为三级及以上数字标题（如1.1.1），如果是则不视为标题
            if re.match(r"^\d+(\.\d+){2,}\s*", text):
                return False
            return True

        # 5. 增强格式检测
        if hasattr(paragraph, "runs") and paragraph.runs:
            # 改进加粗检测（允许80%阈值）
            bold_count = sum(1 for run in paragraph.runs if getattr(run.font, 'bold', False))
            if bold_count / len(paragraph.runs) > 0.8:
                return True
                
            # 改进字体检测（全段落扫描）
            large_font_count = 0
            hei_font_count = 0
            for run in paragraph.runs:
                font = run.font
                if not font: continue
                
                # 黑体检测增强
                if font.name and any(k in font.name for k in ["黑体", "Hei", "Heiti", "SimHei"]):
                    hei_font_count += 1
                    
                # 动态字体检测（增加中位数检测）
                if font.size and font.size.pt > median_size * 1.6:
                    large_font_count += 1
            
            # 阈值优化（50%以上run符合特征）
            if hei_font_count/len(paragraph.runs) > 0.5 or large_font_count/len(paragraph.runs) > 0.5:
                return True

        # 6. 增强文本特征
        # 放宽长度限制 (2-50字符)
        if 2 <= len(text) <= 50:
            # 包含标题常见关键词
            title_keywords = ["摘要", "附录", "参考文献", "目录", "致谢", "章节"]
            if doc_type == "contract":
                title_keywords.extend(["合同", "协议", "甲方", "乙方", "鉴于", "条款"])
            elif doc_type == "policy":
                title_keywords.extend(["制度", "规定", "办法", "细则", "总则", "附则"])
                
            if any(kw in text for kw in title_keywords):
                return True
                
            # 结尾符号检测（标题通常不含逗号/句号）
            if not text.endswith(("。", "，", ".", ",")):
                # 增强大写检测（允许短标题）
                if len(text) >= 2 and text.isupper():
                    return True


        return False
