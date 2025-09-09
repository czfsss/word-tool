from collections.abc import Generator
from typing import Any
import tempfile
import os
import json
import re
from datetime import datetime
from difflib import SequenceMatcher
from dify_plugin.file.file import File
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
import docx

from tools.utils.logger_utils import get_logger
from tools.utils.file_utils import get_meta_data, sanitize_filename


class WordCommentTool(Tool):
    # 获取当前模块的日志记录器
    logger = get_logger(__name__)

    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        # 检查python-docx版本
        self.logger.info(f"当前python-docx版本: {docx.__version__}")

        # 获取上传的Word文件
        word_content = tool_parameters.get("word_content")
        comments_json: str = tool_parameters.get("comments_json", "{}")

        if not word_content:
            self.logger.error("未提供Word文件")
            yield self.create_text_message("请提供Word文件")
            return

        # 检查文件类型
        if not isinstance(word_content, File):
            self.logger.error("无效的文件格式，期望File对象")
            yield self.create_text_message("无效的文件格式，期望File对象")
            return

        # 解析批注JSON
        try:
            comments_data = json.loads(comments_json)

            # 支持新的数组格式：[{"原文1":"批注1","原文2":"批注2"},{...}]
            if isinstance(comments_data, list):
                # 合并所有数组元素中的批注对
                comments_dict = {}
                processed_groups = 0
                total_comments = 0

                for i, comment_group in enumerate(comments_data):
                    if isinstance(comment_group, dict):
                        # 过滤掉空对象和无效的批注对
                        valid_comments = {}
                        for key, value in comment_group.items():
                            # 跳过空键、空值或分隔符
                            if (
                                key
                                and value
                                and str(key).strip() != ""
                                and str(value).strip() != ""
                                and not str(key).startswith("--")
                                and str(key) != "合同原文的问题句"
                            ):
                                valid_comments[key] = value

                        if valid_comments:
                            comments_dict.update(valid_comments)
                            processed_groups += 1
                            total_comments += len(valid_comments)
                            self.logger.info(
                                f"处理第{i+1}个对象组，包含{len(valid_comments)}个有效批注对"
                            )
                        else:
                            self.logger.debug(
                                f"第{i+1}个对象组为空或无有效批注，已跳过"
                            )
                    else:
                        self.logger.warning(
                            f"批注JSON数组中第{i+1}个元素不是对象格式，已跳过"
                        )

                self.logger.info(
                    f"解析数组格式批注JSON成功，处理了 {processed_groups} 个有效对象组，共合并 {len(comments_dict)} 个批注对"
                )

                if len(comments_dict) == 0:
                    self.logger.warning("JSON数组中未找到有效的批注对")
                    yield self.create_text_message(
                        "JSON数组中未找到有效的批注对，请检查数据格式"
                    )
                    return

            elif isinstance(comments_data, dict):
                # 兼容旧的对象格式：{"原文":"批注"}
                # 同样过滤无效的批注对
                comments_dict = {}
                for key, value in comments_data.items():
                    if (
                        key
                        and value
                        and str(key).strip() != ""
                        and str(value).strip() != ""
                        and not str(key).startswith("--")
                        and str(key) != "合同原文的问题句"
                    ):
                        comments_dict[key] = value

                self.logger.info(
                    f"解析对象格式批注JSON成功，共 {len(comments_dict)} 个有效批注对"
                )

                if len(comments_dict) == 0:
                    self.logger.warning("JSON对象中未找到有效的批注对")
                    yield self.create_text_message(
                        "JSON对象中未找到有效的批注对，请检查数据格式"
                    )
                    return
            else:
                self.logger.error("批注JSON格式错误，期望数组或对象格式")
                yield self.create_text_message(
                    "批注JSON格式错误，期望数组格式 [{},...] 或对象格式 {}"
                )
                return

        except json.JSONDecodeError as e:
            self.logger.error(f"JSON解析失败: {str(e)}")
            yield self.create_text_message(f"JSON解析失败: {str(e)}")
            return

        # 获取批注者信息、自定义文件名和相似度阈值
        author = tool_parameters.get("author", "批注者")
        custom_filename = tool_parameters.get("output_filename", "").strip()
        similarity_threshold = tool_parameters.get("similarity_threshold", 0.8)

        # 验证相似度阈值范围
        if not isinstance(similarity_threshold, (int, float)) or not (
            0.1 <= similarity_threshold <= 1.0
        ):
            similarity_threshold = 0.8
            self.logger.warning(f"相似度阈值无效，使用默认值 0.8")

        self.logger.info(
            f"开始处理Word文档批注，文件名: {word_content.filename if word_content.filename else '未知'}，批注数量: {len(comments_dict)}，批注者: {author}，自定义文件名: {custom_filename if custom_filename else '未设置'}，相似度阈值: {similarity_threshold:.2%}"
        )

        try:
            # 创建临时文件保存上传的Word文件
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as temp_input_file:
                # 获取文件内容（字节）并写入临时文件
                temp_input_file.write(word_content.blob)
                temp_input_path = temp_input_file.name

            # 创建临时文件保存处理后的Word文件
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as temp_output_file:
                temp_output_path = temp_output_file.name

            # 调用批注添加函数
            self.logger.info("开始添加批注到Word文档")
            comment_count = self.add_native_comments_to_document(
                temp_input_path,
                temp_output_path,
                comments_dict,
                author,
                similarity_threshold,
            )
            self.logger.info(f"成功添加了 {comment_count} 个批注")

            # 读取处理后的Word文件内容
            with open(temp_output_path, "rb") as docx_file:
                docx_blob = docx_file.read()

            # 处理自定义文件名
            processed_filename = None
            if custom_filename:
                # 清理并处理文件名
                processed_filename = sanitize_filename(custom_filename)
                self.logger.info(f"使用自定义文件名: {processed_filename}")
            else:
                # 使用原始文件名的基本名称（去掉扩展名）
                if word_content.filename:
                    base_name = os.path.splitext(word_content.filename)[0]
                    processed_filename = f"{base_name}_commented"
                else:
                    processed_filename = "commented_document"
                self.logger.info(f"使用默认文件名: {processed_filename}")

            # 使用create_blob_message返回处理后的Word文件
            yield self.create_blob_message(
                blob=docx_blob,
                meta=get_meta_data(
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    output_filename=processed_filename,
                ),
            )

            # 清理临时文件
            os.unlink(temp_input_path)
            os.unlink(temp_output_path)
            self.logger.info("Word批注处理完成，临时文件已清理")

        except Exception as e:
            self.logger.exception("处理Word文档批注时发生异常")
            yield self.create_text_message(f"处理Word文档批注时出错: {str(e)}")

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """
        计算两个文本的相似度

        Args:
            text1: 第一个文本
            text2: 第二个文本

        Returns:
            float: 相似度，范围0-1
        """
        # 清理文本，去除多余空格和标点符号对比
        clean_text1 = re.sub(r"\s+", "", text1.strip())
        clean_text2 = re.sub(r"\s+", "", text2.strip())

        # 使用SequenceMatcher计算相似度
        similarity = SequenceMatcher(None, clean_text1, clean_text2).ratio()
        return similarity

    def _find_fuzzy_match(
        self, target_text: str, paragraph_text: str, threshold: float = 0.8
    ) -> tuple:
        """
        在段落中寻找与目标文本模糊匹配的部分，根据目标文本是否包含换行符采用不同策略

        Args:
            target_text: 目标文本（批注的key）
            paragraph_text: 段落文本
            threshold: 相似度阈值（默认0.8即80%）

        Returns:
            tuple: (是否找到匹配, 匹配的文本, 相似度)
        """
        # 首先尝试精确匹配
        if target_text in paragraph_text:
            return True, target_text, 1.0

        # 判断目标文本是否包含换行符，采用不同的匹配策略
        if "\n" in target_text:
            # 情况1：目标文本包含换行符，表示是多段落组合
            return self._find_multi_paragraph_match(
                target_text, paragraph_text, threshold
            )
        else:
            # 情况2：目标文本不包含换行符，表示是单句或单段落
            return self._find_single_sentence_match(
                target_text, paragraph_text, threshold
            )

    def _find_single_sentence_match(
        self, target_text: str, paragraph_text: str, threshold: float
    ) -> tuple:
        """
        处理单句或单段落的匹配

        Args:
            target_text: 目标文本（不包含换行符）
            paragraph_text: 段落文本
            threshold: 相似度阈值

        Returns:
            tuple: (是否找到匹配, 匹配的文本, 相似度)
        """
        best_match = None
        best_similarity = 0.0

        # 1. 尝试标准化后的精确匹配（去除多余空格和标点）
        normalized_target = re.sub(r"\s+", " ", target_text.strip())
        normalized_paragraph = re.sub(r"\s+", " ", paragraph_text.strip())

        if normalized_target in normalized_paragraph:
            # 找到标准化后的匹配，简化位置映射
            start_idx = normalized_paragraph.find(normalized_target)
            # 直接使用标准化文本的位置（近似定位）
            end_idx = start_idx + len(normalized_target)
            if end_idx <= len(paragraph_text):
                matched_text = paragraph_text[start_idx:end_idx]
            else:
                matched_text = normalized_target  # 降级到标准化文本
            return True, matched_text, 1.0

        # 2. 改进的句子分割策略，支持更多标点符号
        sentences = re.split(r"[。！？.!?；;,、]+", paragraph_text)

        # 3. 尝试单个句子匹配
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 5:  # 降低最小长度阈值
                continue

            similarity = self._calculate_similarity(target_text, sentence)
            if similarity > best_similarity and similarity >= threshold:
                best_similarity = similarity
                best_match = sentence

        # 4. 如果目标文本较长，尝试多句子组合匹配
        if not best_match and len(target_text) > 30:
            # 尝试相邻句子组合
            for i in range(len(sentences) - 1):
                # 尝试两个相邻句子
                combined_sentence = (
                    sentences[i].strip() + " " + sentences[i + 1].strip()
                ).strip()
                if len(combined_sentence) < 10:
                    continue

                similarity = self._calculate_similarity(target_text, combined_sentence)
                if similarity > best_similarity and similarity >= threshold:
                    best_similarity = similarity
                    best_match = combined_sentence

                # 尝试三个相邻句子
                if i < len(sentences) - 2:
                    combined_three = (
                        combined_sentence + " " + sentences[i + 2].strip()
                    ).strip()
                    if len(combined_three) < 15:
                        continue

                    similarity = self._calculate_similarity(target_text, combined_three)
                    if similarity > best_similarity and similarity >= threshold:
                        best_similarity = similarity
                        best_match = combined_three

        # 5. 如果没有找到句子级别的匹配，尝试滑动窗口匹配
        if not best_match and len(target_text) > 20:
            target_len = len(target_text)
            # 动态调整窗口大小
            window_size = max(target_len - 10, target_len // 2, 30)

            # 使用不同的窗口大小进行匹配
            for window_factor in [1.0, 0.8, 0.6, 1.2]:
                current_window_size = int(window_size * window_factor)

                for i in range(len(paragraph_text) - current_window_size + 1):
                    window_text = paragraph_text[i : i + current_window_size]
                    similarity = self._calculate_similarity(target_text, window_text)

                    if similarity > best_similarity and similarity >= threshold:
                        best_similarity = similarity
                        best_match = window_text

                # 如果已经找到匹配，可以提前退出
                if best_match:
                    break

        # 6. 如果仍然没有找到匹配，尝试关键词匹配
        if not best_match:
            # 提取目标文本中的关键词（去除常见停用词）
            target_words = re.findall(r"\b\w+\b", target_text.lower())
            # 简单过滤掉一些常见词
            stop_words = {
                "的",
                "了",
                "和",
                "是",
                "在",
                "我",
                "有",
                "这",
                "个",
                "那",
                "你",
                "会",
                "说",
                "the",
                "a",
                "an",
                "and",
                "or",
                "but",
                "in",
                "on",
                "at",
                "to",
                "for",
                "of",
                "with",
                "by",
                "is",
                "are",
                "was",
                "were",
                "be",
                "been",
                "being",
                "have",
                "has",
                "had",
                "do",
                "does",
                "did",
                "will",
                "would",
                "could",
                "should",
                "may",
                "might",
                "must",
                "can",
                "this",
                "that",
                "these",
                "those",
            }
            key_words = [
                word
                for word in target_words
                if word not in stop_words and len(word) > 1
            ]

            if key_words:
                # 计算段落中包含的关键词比例
                paragraph_words = re.findall(r"\b\w+\b", paragraph_text.lower())
                matched_words = [word for word in key_words if word in paragraph_words]

                if matched_words:
                    keyword_ratio = len(matched_words) / len(key_words)
                    # 如果关键词匹配比例较高，则认为找到匹配
                    if keyword_ratio >= 0.7:  # 70%的关键词匹配
                        # 尝试找到包含最多关键词的文本片段
                        best_match = self._find_best_keyword_match(
                            paragraph_text, key_words
                        )
                        best_similarity = max(
                            threshold, keyword_ratio
                        )  # 使用关键词比例作为相似度

        if best_match:
            return True, best_match, best_similarity
        else:
            return False, None, 0.0

    def _find_multi_paragraph_match(
        self, target_text: str, paragraph_text: str, threshold: float
    ) -> tuple:
        """
        处理多段落组合的匹配 - 注意：这个函数在新设计中仅用于单段落内的多段落文本匹配
        实际的多段落匹配应该在文档级别进行，而不是在单个段落内

        Args:
            target_text: 目标文本（包含换行符）
            paragraph_text: 段落文本
            threshold: 相似度阈值

        Returns:
            tuple: (是否找到匹配, 匹配的文本, 相似度)
        """
        # 1. 将目标文本按换行符分割成多个段落
        target_paragraphs = [p.strip() for p in target_text.split("\n") if p.strip()]

        if not target_paragraphs:
            return False, None, 0.0

        # 2. 如果只有一个段落，回退到单句匹配
        if len(target_paragraphs) == 1:
            return self._find_single_sentence_match(
                target_paragraphs[0], paragraph_text, threshold
            )

        # 3. 新策略：尝试在单个段落中匹配整个目标文本（去除换行符）
        # 将多段落目标文本合并为单行，然后进行匹配
        combined_target = " ".join(target_paragraphs)
        found, matched_text, similarity = self._find_single_sentence_match(
            combined_target, paragraph_text, threshold
        )

        if found:
            return True, matched_text, similarity

        # 4. 如果直接合并匹配失败，尝试部分匹配策略
        # 计算每个目标段落在当前段落中的最佳匹配
        partial_matches = []
        total_similarity = 0
        matched_any = False

        for target_para in target_paragraphs:
            found_para, matched_text_para, similarity_para = (
                self._find_single_sentence_match(
                    target_para,
                    paragraph_text,
                    threshold * 0.7,  # 降低阈值以增加匹配可能性
                )
            )

            if found_para:
                partial_matches.append((matched_text_para, similarity_para))
                total_similarity += similarity_para
                matched_any = True
            else:
                # 即使某个段落没找到，也不立即失败，而是继续尝试其他段落
                self.logger.debug(
                    f"在当前段落中未找到目标子段落: '{target_para[:50]}...'"
                )

        # 5. 如果至少匹配到了一些内容，计算综合得分
        if (
            matched_any and len(partial_matches) >= len(target_paragraphs) * 0.5
        ):  # 至少匹配50%的目标段落
            avg_similarity = total_similarity / len(partial_matches)

            # 如果平均相似度达到调整后的阈值
            if avg_similarity >= threshold * 0.8:  # 降低阈值要求
                # 按照在原文中的顺序排序匹配结果
                sorted_matches = []
                for matched_text_para, similarity_para in partial_matches:
                    pos = paragraph_text.find(matched_text_para)
                    if pos != -1:
                        sorted_matches.append((pos, matched_text_para))

                # 按位置排序
                sorted_matches.sort(key=lambda x: x[0])

                # 组合结果
                combined_match = " ".join([match for pos, match in sorted_matches])
                return True, combined_match, avg_similarity

        return False, None, 0.0

    def _find_cross_paragraph_match(
        self,
        target_text: str,
        all_paragraphs: list,
        start_index: int = 0,
        threshold: float = 0.7,
    ) -> tuple:
        """
        在多个连续段落中查找匹配的文本（新增的文档级别多段落匹配算法）

        Args:
            target_text: 目标文本（包含换行符，表示多段落）
            all_paragraphs: 文档中所有段落的列表 [(paragraph_obj, paragraph_text), ...]
            start_index: 开始搜索的段落索引
            threshold: 相似度阈值

        Returns:
            tuple: (是否找到匹配, 匹配的段落范围(start, end), 匹配的文本列表, 平均相似度)
        """
        # 将目标文本按换行符分割成多个段落
        target_paragraphs = [p.strip() for p in target_text.split("\n") if p.strip()]

        if not target_paragraphs or not all_paragraphs:
            return False, None, [], 0.0

        # 如果只有一个目标段落，使用简化逻辑
        if len(target_paragraphs) == 1:
            for i in range(start_index, len(all_paragraphs)):
                paragraph_obj, paragraph_text = all_paragraphs[i]
                found, matched_text, similarity = self._find_single_sentence_match(
                    target_paragraphs[0], paragraph_text, threshold
                )
                if found:
                    return True, (i, i), [matched_text], similarity
            return False, None, [], 0.0

        # 多段落匹配逻辑
        best_match = None
        best_similarity = 0.0

        # 遍历可能的起始段落
        for start_para_idx in range(
            start_index, len(all_paragraphs) - len(target_paragraphs) + 1
        ):
            # 尝试在从当前起始位置开始的连续段落中匹配所有目标段落
            current_matches = []
            current_similarities = []
            matched_paragraphs = []

            for target_idx, target_para in enumerate(target_paragraphs):
                # 计算当前应该检查的文档段落索引
                doc_para_idx = start_para_idx + target_idx

                # 检查是否超出文档范围
                if doc_para_idx >= len(all_paragraphs):
                    break

                paragraph_obj, paragraph_text = all_paragraphs[doc_para_idx]

                # 在当前段落中查找目标段落
                found, matched_text, similarity = self._find_single_sentence_match(
                    target_para,
                    paragraph_text,
                    threshold * 0.8,  # 稍微降低单个段落的匹配要求
                )

                if found:
                    current_matches.append(matched_text)
                    current_similarities.append(similarity)
                    matched_paragraphs.append(doc_para_idx)
                else:
                    # 当前匹配序列中断，尝试下一个起始位置
                    break

            # 检查是否匹配了所有目标段落
            if len(current_matches) == len(target_paragraphs):
                avg_similarity = sum(current_similarities) / len(current_similarities)

                # 如果平均相似度更高，更新最佳匹配
                if avg_similarity > best_similarity and avg_similarity >= threshold:
                    best_similarity = avg_similarity
                    best_match = {
                        "range": (
                            start_para_idx,
                            start_para_idx + len(target_paragraphs) - 1,
                        ),
                        "matches": current_matches,
                        "paragraph_indices": matched_paragraphs,
                    }

            # 如果匹配了部分段落，也尝试部分匹配策略
            elif len(current_matches) >= len(target_paragraphs) * 0.6:  # 至少匹配60%
                partial_avg_similarity = sum(current_similarities) / len(
                    current_similarities
                )

                # 对部分匹配降低要求
                if (
                    partial_avg_similarity > best_similarity
                    and partial_avg_similarity >= threshold * 0.8
                ):
                    best_similarity = partial_avg_similarity
                    best_match = {
                        "range": (
                            start_para_idx,
                            start_para_idx + len(current_matches) - 1,
                        ),
                        "matches": current_matches,
                        "paragraph_indices": matched_paragraphs,
                    }

        if best_match:
            return True, best_match["range"], best_match["matches"], best_similarity

        return False, None, [], 0.0

    def _find_flexible_cross_paragraph_match(
        self,
        target_text: str,
        all_paragraphs: list,
        start_index: int = 0,
        threshold: float = 0.7,
    ) -> tuple:
        """
        灵活的跨段落匹配算法 - 允许目标段落不严格按顺序匹配

        Args:
            target_text: 目标文本（包含换行符）
            all_paragraphs: 文档中所有段落的列表
            start_index: 开始搜索的段落索引
            threshold: 相似度阈值

        Returns:
            tuple: (是否找到匹配, 匹配的段落索引列表, 匹配的文本列表, 平均相似度)
        """
        target_paragraphs = [p.strip() for p in target_text.split("\n") if p.strip()]

        if not target_paragraphs or not all_paragraphs:
            return False, [], [], 0.0

        # 为每个目标段落在文档中查找最佳匹配
        paragraph_matches = []

        for target_para in target_paragraphs:
            best_match_for_target = None
            best_similarity_for_target = 0.0

            # 在文档的所有段落中查找当前目标段落的最佳匹配
            for doc_idx in range(start_index, len(all_paragraphs)):
                paragraph_obj, paragraph_text = all_paragraphs[doc_idx]

                found, matched_text, similarity = self._find_single_sentence_match(
                    target_para, paragraph_text, threshold * 0.7
                )

                if found and similarity > best_similarity_for_target:
                    best_similarity_for_target = similarity
                    best_match_for_target = {
                        "doc_index": doc_idx,
                        "matched_text": matched_text,
                        "similarity": similarity,
                        "target_paragraph": target_para,
                    }

            if best_match_for_target:
                paragraph_matches.append(best_match_for_target)

        # 检查是否找到了足够的匹配
        if len(paragraph_matches) >= len(target_paragraphs) * 0.6:  # 至少匹配60%
            # 按文档中的位置排序
            paragraph_matches.sort(key=lambda x: x["doc_index"])

            # 计算平均相似度
            avg_similarity = sum(
                match["similarity"] for match in paragraph_matches
            ) / len(paragraph_matches)

            if avg_similarity >= threshold * 0.8:  # 调整阈值要求
                matched_indices = [match["doc_index"] for match in paragraph_matches]
                matched_texts = [match["matched_text"] for match in paragraph_matches]

                return True, matched_indices, matched_texts, avg_similarity

        return False, [], [], 0.0

    def _process_multi_paragraph_comments(
        self, doc, multi_paragraph_comments, author, initials, similarity_threshold=0.7
    ):
        """
        处理多段落批注（新增功能）

        Args:
            doc: Word文档对象
            multi_paragraph_comments: 多段落批注字典
            author: 批注者
            initials: 批注者缩写
            similarity_threshold: 相似度阈值

        Returns:
            int: 成功添加的批注数量
        """
        comment_count = 0

        # 准备所有段落的列表（不包括表格）
        all_paragraphs = [
            (para, para.text.strip()) for para in doc.paragraphs if para.text.strip()
        ]

        self.logger.info(
            f"开始处理 {len(multi_paragraph_comments)} 个多段落批注，文档共有 {len(all_paragraphs)} 个非空段落"
        )

        processed_comments = set()  # 记录已处理的批注，避免重复

        for summary, comment_text in multi_paragraph_comments.items():
            if summary in processed_comments:
                continue

            self.logger.debug(f"处理多段落批注: '{summary[:100]}...'")

            # 使用严格的跨段落匹配算法
            found_strict, range_strict, matches_strict, similarity_strict = (
                self._find_cross_paragraph_match(
                    summary, all_paragraphs, 0, similarity_threshold
                )
            )

            if found_strict:
                self.logger.info(
                    f"使用严格算法找到多段落匹配 (相似度: {similarity_strict:.2%})"
                )

                # 为匹配的段落范围添加批注
                success = self._add_cross_paragraph_comment(
                    doc,
                    all_paragraphs,
                    range_strict,
                    matches_strict,
                    comment_text,
                    author,
                    initials,
                )

                if success:
                    comment_count += 1
                    processed_comments.add(summary)
                    self.logger.info(f"成功为多段落批注添加批注: '{summary[:50]}...'")
                continue

            # 如果严格匹配失败，尝试灵活匹配
            found_flexible, indices_flexible, matches_flexible, similarity_flexible = (
                self._find_flexible_cross_paragraph_match(
                    summary, all_paragraphs, 0, similarity_threshold
                )
            )

            if found_flexible:
                self.logger.info(
                    f"使用灵活算法找到多段落匹配 (相似度: {similarity_flexible:.2%})"
                )

                # 为匹配的段落添加批注（灵活匹配可能不连续）
                success = self._add_flexible_cross_paragraph_comment(
                    doc,
                    all_paragraphs,
                    indices_flexible,
                    matches_flexible,
                    comment_text,
                    author,
                    initials,
                )

                if success:
                    comment_count += 1
                    processed_comments.add(summary)
                    self.logger.info(
                        f"成功为多段落批注添加批注(灵活匹配): '{summary[:50]}...'"
                    )
                continue

            # 如果多段落匹配都失败，记录日志
            self.logger.warning(f"多段落批注未找到匹配: '{summary[:50]}...'")

        self.logger.info(f"多段落批注处理完成，成功添加 {comment_count} 个批注")
        return comment_count

    def _add_cross_paragraph_comment(
        self,
        doc,
        all_paragraphs,
        paragraph_range,
        matched_texts,
        comment_text,
        author,
        initials,
    ):
        """
        为连续的多个段落添加批注 - 改进版：将多个段落合并为一个批注区域

        Args:
            doc: Word文档对象
            all_paragraphs: 所有段落列表
            paragraph_range: 段落范围 (start_index, end_index)
            matched_texts: 匹配的文本列表
            comment_text: 批注内容
            author: 批注者
            initials: 批注者缩写

        Returns:
            bool: 是否成功添加
        """
        try:
            start_idx, end_idx = paragraph_range

            if start_idx >= len(all_paragraphs) or end_idx >= len(all_paragraphs):
                self.logger.error(f"段落范围超出文档范围: {paragraph_range}")
                return False

            # 收集所有需要批注的段落和对应的匹配文本
            target_paragraphs = []
            target_matched_texts = []

            for i in range(start_idx, min(end_idx + 1, len(all_paragraphs))):
                paragraph_obj = all_paragraphs[i][0]
                target_paragraphs.append(paragraph_obj)

                # 获取对应的匹配文本
                if i - start_idx < len(matched_texts):
                    target_matched_texts.append(matched_texts[i - start_idx])
                else:
                    # 如果没有对应的匹配文本，使用段落的完整文本
                    target_matched_texts.append(paragraph_obj.text.strip())

            if not target_paragraphs:
                return False

            # 方案1：尝试使用python-docx的范围批注功能（如果支持）
            success = self._add_multi_paragraph_comment(
                doc,
                target_paragraphs,
                target_matched_texts,
                comment_text,
                author,
                initials,
            )

            if success:
                self.logger.info(f"成功为 {len(target_paragraphs)} 个段落添加批注")
                return True

            # 方案3：降级到第一个段落批注，但在批注内容中说明涵盖的段落范围
            enhanced_comment = (
                f"{comment_text}\n\n[此批注涵盖第{start_idx+1}-{end_idx+1}段的内容]"
            )
            success = self._add_native_comment_to_paragraph(
                doc,
                target_paragraphs[0],
                target_matched_texts[0],
                enhanced_comment,
                author,
                initials,
            )

            if success:
                self.logger.info(f"使用降级方案为多段落添加批注")
                return True

            return False

        except Exception as e:
            self.logger.error(f"添加跨段落批注时出错: {str(e)}")
            return False

    def _add_multi_paragraph_comment(
        self,
        doc,
        target_paragraphs,
        target_matched_texts,
        comment_text,
        author,
        initials,
    ):
        """
        为多个段落创建批注 - 合并原有的两个功能

        Args:
            doc: Word文档对象
            target_paragraphs: 目标段落列表
            target_matched_texts: 匹配的文本列表
            comment_text: 批注内容
            author: 批注者
            initials: 批注者缩写

        Returns:
            bool: 是否成功
        """
        try:
            if not target_paragraphs:
                return False

            # 方案1：尝试使用python-docx的原生多段落批注API
            all_target_runs = []
            for i, paragraph in enumerate(target_paragraphs):
                if i < len(target_matched_texts):
                    matched_text = target_matched_texts[i]
                    paragraph_runs = self._find_runs_for_text(paragraph, matched_text)
                    all_target_runs.extend(paragraph_runs)
                else:
                    all_target_runs.extend(paragraph.runs)

            if all_target_runs:
                try:
                    # 尝试原生多段落批注
                    if len(all_target_runs) == 1:
                        comment = doc.add_comment(
                            runs=all_target_runs[0],
                            text=comment_text,
                            author=author,
                            initials=initials,
                        )
                    else:
                        comment = doc.add_comment(
                            runs=all_target_runs,
                            text=comment_text,
                            author=author,
                            initials=initials,
                        )
                    self.logger.debug(f"成功使用原生范围批注API")
                    return True
                except (AttributeError, TypeError):
                    self.logger.debug(f"原生范围批注API不支持，使用合并文本方案")

            # 方案2：如果原生批注不支持，使用合并文本的单个批注
            merged_text = "\n".join(target_matched_texts)
            first_paragraph = target_paragraphs[0]
            enhanced_comment = f"{comment_text}\n\n[批注范围 - 共{len(target_paragraphs)}个段落]:\n{merged_text[:500]}{'...' if len(merged_text) > 500 else ''}"

            success = self._add_native_comment_to_paragraph(
                doc,
                first_paragraph,
                target_matched_texts[0],
                enhanced_comment,
                author,
                initials,
            )
            return success

        except Exception as e:
            self.logger.error(f"创建多段落批注时出错: {str(e)}")
            return False

    def _find_runs_for_text(self, paragraph, target_text):
        """
        在段落中找到包含目标文本的runs

        Args:
            paragraph: 段落对象
            target_text: 目标文本

        Returns:
            list: 包含目标文本的runs列表
        """
        target_runs = []
        remaining_text = target_text.strip()

        try:
            for run in paragraph.runs:
                if not remaining_text:
                    break

                if run.text and run.text.strip():
                    # 检查run文本是否包含在剩余文本中
                    if run.text.strip() in remaining_text:
                        target_runs.append(run)
                        # 从剩余文本中移除已匹配的部分
                        remaining_text = remaining_text.replace(
                            run.text.strip(), "", 1
                        ).strip()
                    elif remaining_text in run.text:
                        # 如果剩余文本完全包含在当前run中
                        target_runs.append(run)
                        remaining_text = ""
                        break

            # 如果没有找到精确匹配，返回段落的所有runs
            if not target_runs:
                target_runs = [run for run in paragraph.runs if run.text.strip()]

            return target_runs

        except Exception as e:
            self.logger.error(f"查找目标文本的runs时出错: {str(e)}")
            return [run for run in paragraph.runs if run.text.strip()]

    def _add_flexible_cross_paragraph_comment(
        self,
        doc,
        all_paragraphs,
        paragraph_indices,
        matched_texts,
        comment_text,
        author,
        initials,
    ):
        """
        为灵活匹配的多个段落添加批注（可能不连续） - 改进版：合并显示多个段落的内容

        Args:
            doc: Word文档对象
            all_paragraphs: 所有段落列表
            paragraph_indices: 匹配的段落索引列表
            matched_texts: 匹配的文本列表
            comment_text: 批注内容
            author: 批注者
            initials: 批注者缩写

        Returns:
            bool: 是否成功添加
        """
        try:
            if not paragraph_indices or not matched_texts:
                return False

            # 收集所有匹配的段落
            target_paragraphs = []
            for idx in paragraph_indices:
                if idx < len(all_paragraphs):
                    target_paragraphs.append(all_paragraphs[idx][0])

            if not target_paragraphs:
                return False

            # 尝试使用多段落批注
            success = self._add_multi_paragraph_comment(
                doc, target_paragraphs, matched_texts, comment_text, author, initials
            )

            if success:
                self.logger.info(
                    f"成功为灵活匹配的 {len(target_paragraphs)} 个段落添加批注"
                )
                return True

            # 最后降级到第一个段落批注并说明跨段落情况
            first_para_idx = paragraph_indices[0]
            paragraph_list_str = ", ".join(
                [f"第{idx+1}段" for idx in paragraph_indices]
            )
            enhanced_comment = (
                f"{comment_text}\n\n[此批注涵盖多个段落: {paragraph_list_str}]"
            )

            success = self._add_native_comment_to_paragraph(
                doc,
                target_paragraphs[0],
                matched_texts[0],
                enhanced_comment,
                author,
                initials,
            )

            if success:
                self.logger.info(f"使用降级方案为灵活匹配的多段落添加批注")
                return True

            return False

        except Exception as e:
            self.logger.error(f"添加灵活跨段落批注时出错: {str(e)}")
            return False

    def _process_paragraph_comments(
        self, doc, paragraph, comments_dict, author, initials, similarity_threshold=0.8
    ):
        """处理段落中的批注（支持模糊匹配）"""
        comment_count = 0
        paragraph_text = paragraph.text

        # 跳过空段落
        if not paragraph_text.strip():
            return comment_count

        self.logger.debug(f"正在处理段落：{paragraph_text[:100]}...")

        # 统计在当前段落中找到的批注数量
        found_comments = []

        for summary, comment_text in comments_dict.items():
            # 跳过空的或无效的批注
            if not summary or not comment_text:
                continue

            # 使用模糊匹配查找相似的文本
            found, matched_text, similarity = self._find_fuzzy_match(
                summary, paragraph_text, threshold=similarity_threshold
            )

            if found:
                found_comments.append((summary, comment_text, matched_text, similarity))

                if similarity == 1.0:
                    self.logger.debug(f"在段落中精确匹配到摘要: '{summary[:50]}...'")
                else:
                    self.logger.info(
                        f"在段落中模糊匹配到摘要 (相似度: {similarity:.2%}):"
                    )
                    self.logger.info(f"  原摘要: '{summary[:50]}...'")
                    self.logger.info(f"  匹配文本: '{matched_text[:50]}...'")

                try:
                    # 使用匹配到的文本添加批注
                    success = self._add_native_comment_to_paragraph(
                        doc, paragraph, matched_text, comment_text, author, initials
                    )
                    if success:
                        comment_count += 1
                        self.logger.info(
                            f"成功为摘要 '{summary[:30]}...' 添加批注 (相似度: {similarity:.2%})"
                        )
                    else:
                        self.logger.warning(f"为摘要 '{summary[:30]}...' 添加批注失败")
                except Exception as e:
                    self.logger.warning(
                        f"为摘要 '{summary[:30]}...' 添加批注时出错: {str(e)}"
                    )
                    continue
            else:
                # 记录未找到匹配的情况
                self.logger.debug(f"未找到匹配 (相似度 < 80%): '{summary[:50]}...'")

        if found_comments:
            self.logger.debug(
                f"在当前段落中找到 {len(found_comments)} 个匹配的批注，成功添加 {comment_count} 个"
            )

        return comment_count

    def _process_table_comments(
        self, doc, table, comments_dict, author, initials, similarity_threshold=0.8
    ):
        """处理表格中的批注"""
        comment_count = 0

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    comment_count += self._process_paragraph_comments(
                        doc,
                        paragraph,
                        comments_dict,
                        author,
                        initials,
                        similarity_threshold,
                    )

        return comment_count

    def _add_native_comment_to_paragraph(
        self, doc, paragraph, target_text, comment_text, author, initials
    ):
        """
        使用python-docx原生API向段落添加批注

        Args:
            doc: Word文档对象
            paragraph: 段落对象
            target_text: 要批注的目标文本
            comment_text: 批注内容
            author: 批注者
            initials: 批注者缩写

        Returns:
            bool: 是否成功添加批注
        """
        try:
            # 查找所有属于目标文本的Run
            target_runs = []
            remaining_text = target_text

            # 遍历所有Run，找到属于目标文本的部分
            for run in paragraph.runs:
                if not remaining_text:  # 如果已经找到所有目标文本，退出循环
                    break

                if run.text and run.text in remaining_text:
                    # 找到Run文本在剩余文本中的位置
                    pos = remaining_text.find(run.text)
                    if pos == 0:  # 如果Run文本是剩余文本的开头部分
                        target_runs.append(run)
                        # 从剩余文本中移除已匹配的部分
                        remaining_text = remaining_text[len(run.text) :].lstrip()
                    elif pos > 0:  # 如果Run文本在剩余文本中间，可能需要分割Run
                        # 这种情况下，我们需要分割Run以匹配目标文本的开头
                        target_part = remaining_text[: len(remaining_text) - pos]
                        split_run = self._split_run_for_comment(
                            paragraph, run, target_part
                        )
                        if split_run:
                            target_runs.append(split_run)
                            remaining_text = remaining_text[len(target_part) :].lstrip()
                        break
                elif run.text and remaining_text in run.text:
                    # Run文本包含剩余文本，需要分割Run
                    pos = run.text.find(remaining_text)
                    if pos >= 0:
                        # 分割Run以获取精确匹配的部分
                        split_run = self._split_run_for_comment(
                            paragraph, run, remaining_text
                        )
                        if split_run:
                            target_runs.append(split_run)
                            remaining_text = ""  # 已找到所有目标文本
                        break

            if not target_runs:
                self.logger.warning(f"未找到包含文本 '{target_text}' 的Run")
                return False

            # 如果还有剩余文本未匹配，说明目标文本跨越了多个Run或段落
            if remaining_text:
                self.logger.warning(
                    f"目标文本 '{target_text}' 跨越了多个Run或段落，可能无法完整批注"
                )
                # 仍然尝试为已匹配的部分添加批注

            # 使用python-docx 1.2.0的原生批注API
            try:
                if len(target_runs) == 1:
                    comment = doc.add_comment(
                        runs=target_runs[0],
                        text=comment_text,
                        author=author,
                        initials=initials,
                    )
                else:
                    # 如果有多个Run，需要合并它们
                    comment = doc.add_comment(
                        runs=target_runs,
                        text=comment_text,
                        author=author,
                        initials=initials,
                    )
                self.logger.debug(f"成功使用原生API添加批注")
                return True
            except AttributeError:
                # 如果不支持原生批注API，使用备用方案
                self.logger.warning("当前版本不支持原生批注API，使用备用方案")
                # 如果有多个Run，只对第一个Run使用备用方案
                return self._add_fallback_comment(target_runs[0], comment_text, author)

        except Exception as e:
            self.logger.error(f"添加原生批注时出错: {str(e)}")
            return False

    def _split_run_for_comment(self, paragraph, run, target_text):
        """
        分割Run以精确定位批注文本

        Args:
            paragraph: 段落对象
            run: 要分割的Run对象
            target_text: 目标文本

        Returns:
            Run: 包含目标文本的新Run对象
        """
        try:
            # 使用partition将文本拆分为：前缀、目标、后缀
            before, matched, after = run.text.partition(target_text)

            if not matched:
                return None

            # 保存原Run的样式
            original_bold = run.bold
            original_italic = run.italic
            original_underline = run.underline
            original_font_size = run.font.size
            original_font_name = run.font.name

            # 更新原Run为前缀文本
            run.text = before

            # 创建目标文本的新Run
            target_run = paragraph.add_run(matched)
            target_run.bold = original_bold
            target_run.italic = original_italic
            target_run.underline = original_underline
            if original_font_size:
                target_run.font.size = original_font_size
            if original_font_name:
                target_run.font.name = original_font_name

            # 创建后缀文本的新Run
            if after:
                after_run = paragraph.add_run(after)
                after_run.bold = original_bold
                after_run.italic = original_italic
                after_run.underline = original_underline
                if original_font_size:
                    after_run.font.size = original_font_size
                if original_font_name:
                    after_run.font.name = original_font_name

            return target_run

        except Exception as e:
            self.logger.error(f"分割Run时出错: {str(e)}")
            return None

    def _find_best_keyword_match(self, paragraph_text: str, key_words: list) -> str:
        """
        在段落中找到包含最多关键词的文本片段，改进版考虑关键词顺序和密度

        Args:
            paragraph_text: 段落文本
            key_words: 关键词列表

        Returns:
            str: 包含最多关键词的文本片段
        """
        # 将段落按句子分割（不包含换行符，因为这是在单个段落内匹配）
        sentences = re.split(r"[。！？.!?；;,、]+", paragraph_text)

        best_match = ""
        best_score = 0

        # 为每个关键词分配权重（可根据实际需求调整）
        keyword_weights = {word: 1.0 for word in key_words}

        # 计算每个句子的得分
        for sentence in sentences:
            if len(sentence.strip()) < 5:  # 跳过太短的句子
                continue

            # 计算关键词得分
            sentence_lower = sentence.lower()
            sentence_words = re.findall(r"\b\w+\b", sentence_lower)

            # 1. 关键词密度得分
            matched_count = sum(1 for word in key_words if word in sentence_words)
            density_score = matched_count / len(sentence_words) if sentence_words else 0

            # 2. 关键词顺序得分（考虑关键词在句子中的顺序）
            order_score = 0
            if matched_count > 1:
                # 找出所有匹配的关键词及其位置
                matched_keywords = [
                    (word, i)
                    for i, word in enumerate(sentence_words)
                    if word in key_words
                ]
                # 按原始关键词列表中的顺序排序
                matched_keywords.sort(
                    key=lambda x: (
                        key_words.index(x[0]) if x[0] in key_words else float("inf")
                    )
                )
                # 计算顺序得分（位置越接近原始顺序，得分越高）
                for i in range(1, len(matched_keywords)):
                    if matched_keywords[i][1] > matched_keywords[i - 1][1]:
                        order_score += 1
                order_score = (
                    order_score / (len(matched_keywords) - 1)
                    if len(matched_keywords) > 1
                    else 0
                )

            # 3. 关键词距离得分（关键词在句子中越集中，得分越高）
            distance_score = 0
            if matched_count > 1:
                positions = [
                    i for i, word in enumerate(sentence_words) if word in key_words
                ]
                if positions:
                    # 计算关键词位置的标准差，标准差越小，关键词越集中
                    avg_position = sum(positions) / len(positions)
                    variance = sum(
                        (pos - avg_position) ** 2 for pos in positions
                    ) / len(positions)
                    # 将方差转换为得分（方差越小，得分越高）
                    max_variance = len(sentence_words) ** 2 / 4  # 最大可能方差
                    distance_score = (
                        1 - (variance / max_variance) if max_variance > 0 else 1
                    )

            # 综合得分（加权平均）
            combined_score = (
                0.5 * density_score + 0.3 * order_score + 0.2 * distance_score
            )

            if combined_score > best_score or (
                combined_score == best_score and len(sentence) < len(best_match)
            ):
                best_score = combined_score
                best_match = sentence

        # 如果单个句子匹配效果不佳，尝试相邻句子组合
        if best_score < 0.5 and len(sentences) > 1:  # 如果综合得分低于0.5
            for i in range(len(sentences) - 1):
                combined = (
                    sentences[i].strip() + " " + sentences[i + 1].strip()
                ).strip()
                if len(combined) < 10:
                    continue

                # 计算组合句子的得分
                combined_words = re.findall(r"\b\w+\b", combined.lower())
                matched_count = sum(1 for word in key_words if word in combined_words)
                density_score = (
                    matched_count / len(combined_words) if combined_words else 0
                )

                if density_score > best_score:
                    best_score = density_score
                    best_match = combined

        return (
            best_match if best_match else paragraph_text[:100]
        )  # 如果没有找到合适的句子，返回段落前100个字符

    def _add_fallback_comment(self, run, comment_text, author):
        """
        备用批注方案（当不支持原生API时）

        Args:
            run: Run对象
            comment_text: 批注内容
            author: 批注者

        Returns:
            bool: 是否成功添加
        """
        try:
            # 高亮目标文本
            run.font.highlight_color = 7  # 黄色高亮

            # 在文本后添加批注标记
            original_text = run.text
            current_time = datetime.now().strftime("%Y-%m-%d %H:%M")
            run.text = (
                f"{original_text} [批注by {author} {current_time}: {comment_text}]"
            )

            self.logger.debug(f"使用备用方案添加批注")
            return True

        except Exception as e:
            self.logger.error(f"添加备用批注时出错: {str(e)}")
            return False

    def add_native_comments_to_document(
        self,
        input_path: str,
        output_path: str,
        comments_dict: dict,
        author: str = "批注者",
        similarity_threshold: float = 0.8,
    ) -> int:
        """
        向Word文档添加真正的批注（使用python-docx原生批注API，支持模糊匹配和跨段落匹配）

        Args:
            input_path: 输入Word文档路径
            output_path: 输出Word文档路径
            comments_dict: 批注字典，key为摘要文本，value为批注内容
            author: 批注者姓名
            similarity_threshold: 模糊匹配的相似度阈值（0.1-1.0）

        Returns:
            int: 成功添加的批注数量
        """
        doc = Document(input_path)
        comment_count = 0

        # 获取作者缩写（取前两个字符）
        initials = author[:2] if len(author) >= 2 else author

        # 记录所有可用的批注键
        self.logger.info(
            f"开始处理批注，共有 {len(comments_dict)} 个批注对，相似度阈值: {similarity_threshold:.2%}:"
        )

        # 分离单段落和多段落批注
        single_paragraph_comments = {}
        multi_paragraph_comments = {}

        for key, value in comments_dict.items():
            if "\n" in key:
                multi_paragraph_comments[key] = value
            else:
                single_paragraph_comments[key] = value

        self.logger.info(
            f"单段落批注: {len(single_paragraph_comments)} 个，多段落批注: {len(multi_paragraph_comments)} 个"
        )

        # 首先处理多段落批注（新增功能）
        if multi_paragraph_comments:
            comment_count += self._process_multi_paragraph_comments(
                doc, multi_paragraph_comments, author, initials, similarity_threshold
            )

        # 然后处理单段落批注
        processed_paragraphs = 0
        total_tables = 0

        if single_paragraph_comments:
            # 遍历文档中的所有段落
            total_paragraphs = len(doc.paragraphs)

            self.logger.info(f"开始处理 {total_paragraphs} 个段落的单段落批注")

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # 只处理非空段落
                    processed_paragraphs += 1
                    paragraph_comments = self._process_paragraph_comments(
                        doc,
                        paragraph,
                        single_paragraph_comments,
                        author,
                        initials,
                        similarity_threshold,
                    )
                    comment_count += paragraph_comments

                    if paragraph_comments > 0:
                        self.logger.debug(
                            f"段落 {i+1} 添加了 {paragraph_comments} 个批注"
                        )

            # 检查表格中的内容
            total_tables = len(doc.tables)
            if total_tables > 0:
                self.logger.info(f"开始处理 {total_tables} 个表格")

                for i, table in enumerate(doc.tables):
                    table_comments = self._process_table_comments(
                        doc,
                        table,
                        single_paragraph_comments,
                        author,
                        initials,
                        similarity_threshold,
                    )
                    comment_count += table_comments

                    if table_comments > 0:
                        self.logger.debug(f"表格 {i+1} 添加了 {table_comments} 个批注")

        self.logger.info(
            f"文档处理完成：处理了 {processed_paragraphs} 个段落和 {total_tables} 个表格，成功添加 {comment_count} 个批注"
        )

        # 保存文档
        doc.save(output_path)
        return comment_count
