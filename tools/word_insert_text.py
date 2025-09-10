from collections.abc import Generator
from typing import Any
import tempfile
import os
import json
from dify_plugin.file.file import File
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tools.utils.logger_utils import get_logger
from tools.utils.file_utils import get_meta_data, sanitize_filename
import markdown
from bs4 import BeautifulSoup


class WordInsertTextTool(Tool):
    # 获取当前模块的日志记录器
    logger = get_logger(__name__)

    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        # 获取上传的Word文件
        word_content = tool_parameters.get("word_content")
        text_to_insert = tool_parameters.get("text_to_insert", "")
        insert_position = tool_parameters.get(
            "insert_position", "end"
        )  # "start" 或 "end"
        font_name = tool_parameters.get("font_name", "宋体")
        font_size = tool_parameters.get("font_size", 12)
        font_color = tool_parameters.get("font_color", "#000000")  # 默认黑色
        custom_filename = tool_parameters.get("output_filename", "").strip()
        is_markdown = tool_parameters.get("is_markdown", False)  # 是否为Markdown文本

        # 验证必需参数
        if not word_content:
            self.logger.error("未提供Word文件")
            yield self.create_text_message("请提供Word文件")
            return

        if not text_to_insert:
            self.logger.error("未提供要插入的文本")
            yield self.create_text_message("请提供要插入的文本")
            return

        # 检查文件类型
        if not isinstance(word_content, File):
            self.logger.error("无效的文件格式，期望File对象")
            yield self.create_text_message("无效的文件格式，期望File对象")
            return

        # 验证插入位置
        if insert_position not in ["start", "end"]:
            self.logger.warning(f"无效的插入位置: {insert_position}，使用默认值 'end'")
            insert_position = "end"

        # 验证字体大小
        try:
            font_size = int(font_size)
            if font_size <= 0:
                font_size = 12
                self.logger.warning("字体大小必须大于0，使用默认值 12")
        except (ValueError, TypeError):
            font_size = 12
            self.logger.warning("无效的字体大小，使用默认值 12")

        # 验证字体颜色格式并转换为十六进制值
        if not self._is_valid_color(font_color):
            self.logger.warning(f"无效的字体颜色: {font_color}，使用默认值 '黑色'")
            font_color = "黑色"

        # 将颜色名称转换为十六进制值
        font_color_hex = self._get_color_hex(font_color)

        self.logger.info(
            f"开始处理Word文档文本插入，文件名: {word_content.filename if word_content.filename else '未知'}，"
            f"插入位置: {insert_position}，字体: {font_name}，大小: {font_size}，颜色: {font_color}，"
            f"是否为Markdown: {is_markdown}"
        )

        try:
            # 创建临时文件保存上传的Word文件
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as temp_input_file:
                # 获取文件内容（字节）并写入临时文件
                temp_input_file.write(word_content.blob)
                temp_input_path = temp_input_file.name

            # 创建临时文件保存处理后的Word文件
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as temp_output_file:
                temp_output_path = temp_output_file.name

            # 调用文本插入函数
            self.logger.info("开始插入文本到Word文档")
            self.insert_text_to_document(
                temp_input_path,
                temp_output_path,
                text_to_insert,
                insert_position,
                font_name,
                font_size,
                font_color_hex,
                is_markdown,
            )
            self.logger.info("成功插入文本到Word文档")

            # 读取处理后的Word文件内容
            with open(temp_output_path, "rb") as docx_file:
                docx_blob = docx_file.read()

            # 处理自定义文件名
            processed_filename = None
            if custom_filename:
                # 清理并处理文件名
                processed_filename = sanitize_filename(custom_filename)
                self.logger.info(f"使用自定义文件名: {processed_filename}")
            else:
                # 使用原始文件名的基本名称（去掉扩展名）
                if word_content.filename:
                    base_name = os.path.splitext(word_content.filename)[0]
                    processed_filename = f"{base_name}_with_text"
                else:
                    processed_filename = "document_with_text"
                self.logger.info(f"使用默认文件名: {processed_filename}")

            # 使用create_blob_message返回处理后的Word文件
            yield self.create_blob_message(
                blob=docx_blob,
                meta=get_meta_data(
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    output_filename=processed_filename,
                ),
            )

            # 清理临时文件
            os.unlink(temp_input_path)
            os.unlink(temp_output_path)

        except Exception as e:
            self.logger.exception(f"处理Word文档时发生异常: {str(e)}")
            yield self.create_text_message(f"处理Word文档时出错: {str(e)}")

    def _get_color_hex(self, color_name_or_hex):
        """将颜色名称或十六进制值转换为十六进制颜色值"""
        # 常见颜色名称到十六进制值的映射
        color_map = {
            "黑色": "#000000",
            "白色": "#FFFFFF",
            "红色": "#FF0000",
            "绿色": "#00FF00",
            "蓝色": "#0000FF",
            "黄色": "#FFFF00",
            "紫色": "#800080",
            "橙色": "#FFA500",
            "粉色": "#FFC0CB",
            "棕色": "#A52A2A",
            "灰色": "#808080",
            "深蓝色": "#00008B",
            "深绿色": "#006400",
            "深红色": "#8B0000",
            "青色": "#00FFFF",
            "金色": "#FFD700",
            "银色": "#C0C0C0",
        }

        # 如果是颜色名称，返回对应的十六进制值
        if color_name_or_hex in color_map:
            return color_map[color_name_or_hex]

        # 如果是十六进制格式，直接返回
        if (
            color_name_or_hex.startswith("#")
            and len(color_name_or_hex) == 7
            and all(c in "0123456789ABCDEFabcdef" for c in color_name_or_hex[1:])
        ):
            return color_name_or_hex

        # 无效的颜色，返回默认黑色
        return "#000000"

    def _is_valid_color(self, color_str):
        """验证颜色字符串是否为有效的颜色名称或十六进制颜色格式"""
        if not isinstance(color_str, str):
            return False

        # 常见颜色名称到十六进制值的映射
        color_map = {
            "黑色",
            "白色",
            "红色",
            "绿色",
            "蓝色",
            "黄色",
            "紫色",
            "橙色",
            "粉色",
            "棕色",
            "灰色",
            "深蓝色",
            "深绿色",
            "深红色",
            "青色",
            "金色",
            "银色",
        }

        # 检查是否为颜色名称
        if color_str in color_map:
            return True

        # 检查是否为十六进制格式
        return (
            color_str.startswith("#")
            and len(color_str) == 7
            and all(c in "0123456789ABCDEFabcdef" for c in color_str[1:])
        )

    def _hex_to_rgb(self, hex_color):
        """将十六进制颜色转换为RGB元组"""
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))

    def _markdown_to_docx(
        self, markdown_text, doc, paragraph, font_name, font_size, font_color
    ):
        """将Markdown文本转换为Word格式并插入到文档中"""
        try:
            # 如果paragraph为None，表示在文档末尾插入，创建第一个段落
            if paragraph is None:
                paragraph = doc.add_paragraph()
                
            # 将Markdown转换为HTML
            html = markdown.markdown(markdown_text)

            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html, "html.parser")

            # 处理不同的HTML元素
            for element in soup.children:
                if element.name is None:  # 文本节点
                    if element.strip():  # 忽略空白文本
                        # 创建新段落
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(element.strip())
                        self._apply_font_format(run, font_name, font_size, font_color)
                elif element.name == "p":
                    # 段落 - 创建新段落
                    paragraph = doc.add_paragraph()
                    for child in element.children:
                        if child.name is None:  # 文本节点
                            if child.strip():  # 忽略空白文本
                                run = paragraph.add_run(child.strip())
                                self._apply_font_format(
                                    run, font_name, font_size, font_color
                                )
                        elif child.name == "strong" or child.name == "b":
                            # 粗体
                            run = paragraph.add_run(child.get_text())
                            run.bold = True
                            self._apply_font_format(
                                run, font_name, font_size, font_color
                            )
                        elif child.name == "em" or child.name == "i":
                            # 斜体
                            run = paragraph.add_run(child.get_text())
                            run.italic = True
                            self._apply_font_format(
                                run, font_name, font_size, font_color
                            )
                        elif child.name == "code":
                            # 行内代码
                            run = paragraph.add_run(child.get_text())
                            self._apply_font_format(
                                run, "Consolas", font_size, font_color
                            )
                        elif child.name == "a":
                            # 链接
                            run = paragraph.add_run(child.get_text())
                            run.underline = True
                            self._apply_font_format(
                                run, font_name, font_size, "#0000FF"
                            )  # 链接默认蓝色
                        else:
                            # 其他标签，直接添加文本
                            run = paragraph.add_run(child.get_text())
                            self._apply_font_format(
                                run, font_name, font_size, font_color
                            )
                elif element.name == "h1":
                    # 标题1 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                    self._apply_font_format(
                        run, font_name, int(font_size * 1.5), font_color
                    )
                elif element.name == "h2":
                    # 标题2 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                    self._apply_font_format(
                        run, font_name, int(font_size * 1.3), font_color
                    )
                elif element.name == "h3":
                    # 标题3 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                    self._apply_font_format(
                        run, font_name, int(font_size * 1.2), font_color
                    )
                elif (
                    element.name == "h4" or element.name == "h5" or element.name == "h6"
                ):
                    # 标题4-6 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                    self._apply_font_format(
                        run, font_name, int(font_size * 1.1), font_color
                    )
                elif element.name == "ul":
                    # 无序列表 - 添加新段落
                    for li in element.find_all("li", recursive=False):
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run("• " + li.get_text())
                        self._apply_font_format(run, font_name, font_size, font_color)
                elif element.name == "ol":
                    # 有序列表 - 添加新段落
                    for i, li in enumerate(element.find_all("li", recursive=False)):
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(f"{i+1}. " + li.get_text())
                        self._apply_font_format(run, font_name, font_size, font_color)
                elif element.name == "blockquote":
                    # 引用块 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    self._apply_font_format(run, font_name, font_size, font_color)
                    paragraph.paragraph_format.left_indent = Pt(18)
                elif element.name == "pre":
                    # 代码块 - 创建新段落
                    paragraph = doc.add_paragraph()
                    code_text = element.get_text()
                    run = paragraph.add_run(code_text)
                    self._apply_font_format(run, "Consolas", font_size, font_color)
                elif element.name == "hr":
                    # 分割线 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run("-" * 50)
                    self._apply_font_format(run, font_name, font_size, font_color)
                else:
                    # 其他标签 - 创建新段落
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(element.get_text())
                    self._apply_font_format(run, font_name, font_size, font_color)

        except Exception as e:
            self.logger.warning(f"Markdown转换时出错: {str(e)}，使用纯文本插入")
            # 如果转换失败，使用纯文本插入
            run = paragraph.add_run(markdown_text)
            self._apply_font_format(run, font_name, font_size, font_color)

    def _apply_font_format(self, run, font_name, font_size, font_color):
        """应用字体格式到文本运行"""
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            r, g, b = self._hex_to_rgb(font_color)
            run.font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            self.logger.warning(f"设置字体格式时出错: {str(e)}")

    def insert_text_to_document(
        self,
        input_path,
        output_path,
        text_to_insert,
        insert_position,
        font_name,
        font_size,
        font_color,
        is_markdown=False,
    ):
        """将文本插入到Word文档的指定位置，并应用指定的字体格式"""
        # 打开Word文档
        doc = Document(input_path)

        # 根据是否为Markdown文本选择不同的插入方式
        if is_markdown:
            # 使用Markdown转换功能
            if insert_position == "start":
                # 记录当前文档段落数量
                original_paragraph_count = len(doc.paragraphs)
                # 在文档末尾插入所有Markdown内容
                self._markdown_to_docx(
                    text_to_insert, doc, None, font_name, font_size, font_color
                )
                
                # 获取所有新增的段落
                new_paragraphs = doc.paragraphs[original_paragraph_count:]
                
                # 将新增段落批量移动到文档开头
                for i, paragraph in enumerate(new_paragraphs):
                    element = paragraph._element
                    doc.element.body.insert(i, element)
            else:
                # 在文档结尾插入，不需要预先创建段落
                self._markdown_to_docx(
                    text_to_insert, doc, None, font_name, font_size, font_color
                )
        else:
            # 使用原来的纯文本插入方式
            # 创建新段落
            if insert_position == "start":
                # 在文档开头插入
                paragraph = doc.add_paragraph()
                # 将段落移动到文档开头
                element = paragraph._element
                doc.element.body.insert(0, element)
            else:
                # 在文档结尾插入
                paragraph = doc.add_paragraph()
                
            run = paragraph.add_run(text_to_insert)
            # 设置字体格式
            try:
                self._apply_font_format(run, font_name, font_size, font_color)
                # 设置段落对齐方式（左对齐）
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception as e:
                self.logger.warning(f"设置字体格式时出错: {str(e)}，使用默认格式")

        # 保存文档
        doc.save(output_path)
