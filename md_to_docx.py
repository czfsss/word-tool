# 导入 io 模块，用于处理字节流
import io
# 导入 logging 模块，用于记录日志
import logging
# 导入 re 模块，用于正则表达式操作
import re
# 从 typing 模块导入 Generator，用于类型提示
from typing import Generator

# 导入 markdown 模块，用于将 Markdown 文本转换为 HTML
import markdown
# 从 dify_plugin 模块导入 Tool 类
from dify_plugin import Tool
# 从 dify_plugin.entities.tool 模块导入 ToolInvokeMessage 类
from dify_plugin.entities.tool import ToolInvokeMessage
# 从 docx 模块导入 Document 类，用于创建和操作 Word 文档
from docx import Document
# 从 docx.oxml.ns 模块导入 qn 函数，用于处理 XML 命名空间
from docx.oxml.ns import qn
# 从 docx.table 模块导入 Table 和 _Cell 类，用于处理表格
from docx.table import Table, _Cell
# 从 docx.text.paragraph 模块导入 Paragraph 类，用于处理段落
from docx.text.paragraph import Paragraph
# 从 docx.text.run 模块导入 Run 类，用于处理文本运行
from docx.text.run import Run
# 从 htmldocx 模块导入 HtmlToDocx 类，用于将 HTML 转换为 Word 文档
from htmldocx import HtmlToDocx

# 从 tools.md_to_docx.font_enum 模块导入 DocxFontEnum 类，用于字体枚举
from tools.md_to_docx.font_enum import DocxFontEnum
# 从 tools.utils.file_utils 模块导入 get_meta_data 函数，用于获取文件元数据
from tools.utils.file_utils import get_meta_data
# 从 tools.utils.mimetype_utils 模块导入 MimeType 类，用于处理文件类型
from tools.utils.mimetype_utils import MimeType
# 从 tools.utils.param_utils 模块导入 get_md_text 函数，用于获取 Markdown 文本
from tools.utils.param_utils import get_md_text


class MarkdownToDocxTool(Tool):
    # 获取当前模块的日志记录器
    logger = logging.getLogger(__name__)

    def _invoke(self, tool_parameters: dict) -> Generator[ToolInvokeMessage, None, None]:
        """
        调用工具
        """
        # 获取参数中的 Markdown 文本，并去除首尾空白
        md_text = get_md_text(tool_parameters, is_strip_wrapper=True)
        try:
            # 旧方法：使用 markdowntodocx 库
            # with NamedTemporaryFile(suffix=".docx", delete=True) as temp_docx_file:
            #     markdownconverter.markdownToWordFromString(string=md_text, outfile=temp_docx_file)
            #     result_file_bytes = Path(temp_docx_file.name).read_bytes()

            # 将 Markdown 文本转换为 HTML，使用 extra 和 toc 扩展
            html = markdown.markdown(text=md_text, extensions=["extra", "toc"])

            # 将 HTML 转换为 Word 文档
            new_parser = HtmlToDocx()
            doc: Document = new_parser.parse_html_string(html)

            # 为文档中所有文本元素设置字体
            try:
                self.set_fonts_for_all_runs(doc)
            except Exception as e:
                self.logger.exception(e)

            # 创建一个字节流对象
            result_bytes_io = io.BytesIO()
            # 将文档保存到字节流中
            doc.save(result_bytes_io)
            # 从字节流中获取文档的字节数据
            result_file_bytes = result_bytes_io.getvalue()
        except Exception as e:
            # 记录文件转换失败的日志
            self.logger.exception("文件转换失败")
            # 生成一个文本消息，提示转换失败并包含错误信息
            yield self.create_text_message(f"将 Markdown 文本转换为 DOCX 文件失败，错误信息: {str(e)}")
            return

        # 生成一个二进制消息，包含转换后的文件字节数据和元数据
        yield self.create_blob_message(
            blob=result_file_bytes,
            meta=get_meta_data(
                mime_type=MimeType.DOCX,
                output_filename=tool_parameters.get("output_filename"),
            ),
        )
        return

    def is_contains_chinese_chars(self, text: str) -> bool:
        # 判断文本中是否包含中文字符
        return bool(re.search(r'[\u4e00-\u9fff]', text))

    def set_chinese_fonts(self, doc):
        # 全局设置字体
        # https://github.com/python-openxml/python-docx/issues/346#issuecomment-1698885586
        # https://zhuanlan.zhihu.com/p/548039429
        # 获取文档的默认样式
        style = doc.styles['Normal']
        # 获取样式的字体对象
        font = style.font
        # 设置默认字体为 Times New Roman
        font.name = DocxFontEnum.TIMES_NEW_ROMAN
        # 获取或添加样式的 rPr 元素
        rPr = style.element.get_or_add_rPr()
        # 设置东亚字体为宋体
        rPr.rFonts.set(qn('w:eastAsia'), DocxFontEnum.SONG_TI)

    def set_fonts_for_all_runs(self, doc: Document):
        """为文档中所有文本元素的英文文本设置 Times New Roman 字体，中文文本设置宋体字体。"""

        # 处理文档中的所有段落
        paragraph: Paragraph
        for paragraph in doc.paragraphs:
            run: Run
            for run in paragraph.runs:
                self.apply_fonts_to_run(run)

        # 处理文档中所有表格中的段落
        table: Table
        for table in doc.tables:
            for row in table.rows:
                cell: _Cell
                for cell in row.cells:
                    paragraph: Paragraph
                    for paragraph in cell.paragraphs:
                        run: Run
                        for run in paragraph.runs:
                            self.apply_fonts_to_run(run)

    def apply_fonts_to_run(self, run: Run):
        # 跳过没有文本的元素
        if not run or not run.text:  
            return

        try:
            # 设置默认字体为 Times New Roman
            run.font.name = DocxFontEnum.TIMES_NEW_ROMAN
            # 设置东亚字体为宋体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), DocxFontEnum.SONG_TI)
            # 设置 ASCII 字体为 Times New Roman
            run._element.rPr.rFonts.set(qn('w:ascii'), DocxFontEnum.TIMES_NEW_ROMAN)
            # 设置高 ANSI 字体为 Times New Roman
            run._element.rPr.rFonts.set(qn('w:hAnsi'), DocxFontEnum.TIMES_NEW_ROMAN)
        except Exception as e:
            # 记录为文本运行应用字体失败的日志
            self.logger.exception("为文本运行应用字体失败")
